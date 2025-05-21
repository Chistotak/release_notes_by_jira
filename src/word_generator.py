# src/word_generator.py
import logging
import re
from typing import Optional, List, Dict, Union  # Добавляем Union для DocxDocument | None

# Используем try-except для импорта get_correct_path, чтобы модуль мог тестироваться автономно,
# но в основном приложении ожидается, что он будет импортирован из src.config_loader.
try:
    from src.config_loader import get_correct_path
except ImportError:
    # Фоллбэк-реализация get_correct_path для автономного тестирования этого модуля.
    # В реальном приложении PyInstaller этот фоллбэк может работать иначе,
    # поэтому важно, чтобы основной импорт из src.config_loader был доступен.
    import sys
    from pathlib import Path


    def get_correct_path(relative_path_str: str) -> Path:
        if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
            return Path(sys._MEIPASS) / relative_path_str
        return Path(__file__).resolve().parent.parent / relative_path_str


    # Если логгер не был настроен глобально, инициализируем его здесь для автономного запуска
    if not logging.getLogger(__name__).hasHandlers():
        logging.basicConfig(level=logging.DEBUG)  # или INFO
    logger_fallback = logging.getLogger(__name__)
    logger_fallback.warning(
        "Не удалось импортировать get_correct_path из src.config_loader. "
        "Используется локальный фоллбэк get_correct_path. "
        "Это может быть нормально при автономном тестировании этого модуля."
    )

from docx import Document
from docx.document import Document as DocxDocument  # Явный импорт для аннотации типа
from docx.shared import Pt, Inches

# from docx.enum.text import WD_ALIGN_PARAGRAPH # Если понадобится

logger = logging.getLogger(__name__)

# Имена стилей Word (могут быть переопределены в config.yaml)
STYLE_NORMAL = 'Normal'
STYLE_HEADING_1 = 'Heading 1'
STYLE_HEADING_2 = 'Heading 2'
STYLE_HEADING_3 = 'Heading 3'
STYLE_HEADING_4 = 'Heading 4'
STYLE_LIST_BULLET = 'List Bullet'
STYLE_TABLE_DEFAULT = 'Table Grid'


def _format_template_string(template_str: str, data_dict: Dict) -> str:
    """
    Заменяет плейсхолдеры вида {ключ} в строке-шаблоне значениями из словаря data_dict.
    Если плейсхолдер не найден или его значение None, он заменяется на пустую строку.
    """

    def replace_match(match_obj):
        key_in_placeholder = match_obj.group(1)
        value_from_data = data_dict.get(key_in_placeholder)
        return str(value_from_data) if value_from_data is not None else ""

    return re.sub(r"\{([\w_.-]+)\}", replace_match, template_str)


def _add_heading_styled(document: DocxDocument, text: Optional[str], level: int, style_to_apply: str):
    """
    Добавляет заголовок в документ, используя указанный уровень
    и применяя указанный стиль.
    """
    if not text or not text.strip():
        logger.debug("Пропуск добавления пустого заголовка в Word.")
        return

    doc_level = max(1, min(4, level))  # Уровни 1-4 для Heading 1-4

    try:
        heading_paragraph = document.add_heading(text='', level=doc_level)
        heading_paragraph.text = text.strip()

        if style_to_apply:
            try:
                current_style_name = ""
                # Проверка существования стиля перед его применением и его имени
                if heading_paragraph.style and hasattr(heading_paragraph.style, 'name'):
                    current_style_name = heading_paragraph.style.name

                # Применяем стиль, если он отличается от уже установленного по умолчанию для уровня,
                # или если это не стандартный "Heading X", а кастомный.
                # Либо можно просто всегда применять style_to_apply, если он задан.
                if current_style_name != style_to_apply:  # Упрощенное условие: всегда пытаемся применить, если задан
                    heading_paragraph.style = style_to_apply
                    logger.debug(f"К заголовку '{text[:30]}...' применен стиль '{style_to_apply}'.")
            except KeyError:
                logger.warning(
                    f"Стиль Word '{style_to_apply}' не найден для заголовка '{text[:30]}...'. Использован стандартный стиль для уровня {doc_level}.")
            except AttributeError:
                logger.warning(
                    f"Не удалось проверить/установить стиль '{style_to_apply}' для заголовка '{text[:30]}...'.")
            except Exception as e_style:
                logger.warning(f"Не удалось применить стиль '{style_to_apply}' к заголовку '{text[:30]}...': {e_style}")
    except Exception as e:
        logger.error(f"Ошибка при добавлении заголовка Word '{text[:30]}...': {e}", exc_info=True)


def _add_task_entry_to_document(
        document: DocxDocument,
        issue_template_str: str,
        task_data: Dict,
        first_line_style_name: str,
        subsequent_line_style_name: str,
        subsequent_line_indent: Optional[Pt] = None
):
    """
    Добавляет запись о задаче в Word документ, форматируя ее по шаблону.
    Каждая строка из отформатированного шаблона становится новым параграфом.
    """
    formatted_entry = _format_template_string(issue_template_str, task_data)
    # Разделяем на строки и удаляем пустые строки, которые могли образоваться из-за пустых плейсхолдеров
    lines = [line.strip() for line in formatted_entry.strip().splitlines() if line.strip()]

    if not lines:
        logger.debug(f"Задача {task_data.get('key', 'UKNOWN_KEY')} не дала контента по шаблону для Word.")
        return

    for i, line_text in enumerate(lines):
        if i == 0:  # Первая строка
            p = document.add_paragraph(line_text, style=first_line_style_name)
        else:  # Последующие строки
            p = document.add_paragraph(line_text, style=subsequent_line_style_name)
            if subsequent_line_style_name == STYLE_NORMAL and subsequent_line_indent is not None:
                try:
                    p.paragraph_format.left_indent = subsequent_line_indent
                except Exception as e:
                    logger.warning(
                        f"Не удалось применить отступ {subsequent_line_indent} к параграфу для задачи {task_data.get('key')}: {e}")


def generate_word_document(processed_data: Dict, app_config: Dict) -> Optional[DocxDocument]:
    """
    Генерирует документ Word (.docx) на основе обработанных данных и конфигурации.

    Args:
        processed_data (Dict): Словарь с данными от data_processor.
        app_config (Dict): Полная конфигурация приложения.

    Returns:
        Optional[DocxDocument]: Объект документа Word или None, если генерация отключена/не удалась.
    """
    logger.info("Начало генерации Word (.docx) документа...")

    word_config = app_config.get('output_formats', {}).get('word', {})
    if not word_config or not word_config.get('enabled', False):
        logger.info("Генерация Word документа отключена в конфигурации или секция 'word' отсутствует.")
        return None

    template_file_path_str = word_config.get('template_path')
    document_obj: DocxDocument  # Аннотация для ясности
    if template_file_path_str:
        actual_template_file_path = get_correct_path(template_file_path_str)
        logger.info(f"Попытка использовать Word шаблон: {actual_template_file_path}")
        try:
            if actual_template_file_path.is_file():
                document_obj = Document(str(actual_template_file_path))
                logger.info(f"Успешно использован Word шаблон: {actual_template_file_path}")
            else:
                logger.warning(f"Файл шаблона Word не найден: {actual_template_file_path}. Создается пустой документ.")
                document_obj = Document()
        except Exception as e:
            logger.warning(
                f"Не удалось загрузить Word шаблон '{actual_template_file_path}': {e}. Создается пустой документ.",
                exc_info=True)
            document_obj = Document()
    else:
        document_obj = Document()
        logger.info("Word шаблон не указан, создается документ со стилями по умолчанию.")

    styles_map = word_config.get('styles', {})
    style_h1 = styles_map.get('main_title', STYLE_HEADING_1)
    style_h2_table = styles_map.get('table_title', STYLE_HEADING_2)
    style_h2_section = styles_map.get('section_title', STYLE_HEADING_2)
    style_h3_ms = styles_map.get('microservice_group', STYLE_HEADING_3)
    style_h4_type = styles_map.get('issue_type_group', STYLE_HEADING_4)
    style_task_first = styles_map.get('list_bullet_first_line', STYLE_LIST_BULLET)
    style_task_multiline = styles_map.get('list_bullet_multiline_indent', STYLE_NORMAL)
    style_table = styles_map.get('table_style', STYLE_TABLE_DEFAULT)

    multiline_indent_val = Pt(20)  # Значение по умолчанию для отступа

    rn_cfg_data = app_config.get('release_notes', {})

    # 1. Главный заголовок
    gv_text = processed_data.get("global_version", "N/A")
    date_text = processed_data.get("current_date", "N/A")
    title_template_str = rn_cfg_data.get('title_template', "Release Notes - {global_version} - {current_date}")
    main_title_str = _format_template_string(title_template_str, {"global_version": gv_text, "current_date": date_text})
    _add_heading_styled(document_obj, main_title_str, level=1, style_to_apply=style_h1)

    # 2. Таблица микросервисов
    ms_table_cfg_data = rn_cfg_data.get('microservices_table', {})
    ms_summary_data_list: List[Dict] = processed_data.get("microservices_summary", [])
    if ms_table_cfg_data.get('enabled', True) and ms_summary_data_list:
        table_title_heading = ms_table_cfg_data.get('title')
        _add_heading_styled(document_obj, table_title_heading, level=2, style_to_apply=style_h2_table)

        cols_config_list: List[Dict] = ms_table_cfg_data.get('columns', [])
        table_col_headers: List[str] = [col.get('header', '') for col in cols_config_list]

        if table_col_headers and any(h.strip() for h in table_col_headers):
            try:
                # Проверяем, достаточно ли данных для строк перед созданием таблицы
                table_content_rows_data = []
                for ms_item in ms_summary_data_list:
                    row = [_format_template_string(col_cfg_item.get('value_placeholder', ''), ms_item)
                           for col_cfg_item in cols_config_list]
                    table_content_rows_data.append(row)

                if table_content_rows_data:  # Создаем таблицу только если есть что в нее положить
                    created_table = document_obj.add_table(rows=1, cols=len(table_col_headers))
                    created_table.style = style_table
                    header_row_cells = created_table.rows[0].cells
                    for i, header_name in enumerate(table_col_headers): header_row_cells[i].text = header_name
                    for data_row in table_content_rows_data:
                        row_cells = created_table.add_row().cells
                        for i, cell_content_str in enumerate(data_row):
                            if i < len(row_cells):  # Защита от несоответствия кол-ва колонок
                                row_cells[i].text = cell_content_str
                    document_obj.add_paragraph()
                else:
                    logger.info("Word: Таблица МС - нет данных для строк после форматирования.")
            except Exception as e:
                logger.error(f"Ошибка при создании таблицы МС в Word: {e}", exc_info=True)
        elif table_col_headers:
            logger.info("Word: Таблица МС - нет данных (ms_summary_data_list пуст).")

    # 3. Информационные секции
    sections_data_map = processed_data.get("sections_data", {})
    sections_meta_map = rn_cfg_data.get('sections', {})

    for section_id, section_meta_cfg in sections_meta_map.items():
        current_section_data = sections_data_map.get(section_id)
        if not current_section_data:
            logger.debug(f"Секция '{section_id}' пропущена в Word (нет данных от data_processor).")
            continue

        _add_heading_styled(document_obj, current_section_data.get('title'), level=2, style_to_apply=style_h2_section)

        issue_template = section_meta_cfg.get('issue_display_template')
        if not issue_template:
            logger.warning(f"Для секции '{section_id}' в Word отсутствует 'issue_display_template'.")
            document_obj.add_paragraph("* Конфигурация отображения задач отсутствует.*", style=style_task_first)
            document_obj.add_paragraph()
            continue

        is_flat_mode = current_section_data.get("disable_grouping", False)

        if is_flat_mode:
            flat_task_list: List[Dict] = current_section_data.get("tasks_flat_list", [])
            if not flat_task_list:
                p = document_obj.add_paragraph(style=style_task_first)
                p.add_run("* Нет задач для отображения в этой секции.*")
            else:
                for task_dict_item in sorted(flat_task_list, key=lambda t: str(t.get("key", ""))):
                    _add_task_entry_to_document(document_obj, issue_template, task_dict_item,
                                                style_task_first, style_task_multiline, multiline_indent_val)
            document_obj.add_paragraph()
        else:
            ms_map = current_section_data.get('microservices', {})
            if not ms_map:
                logger.debug(f"В секции '{section_id}' нет МС с задачами для Word (группировка).")
                continue
            for ms_name_val in sorted(ms_map.keys()):
                ms_render_data = ms_map[ms_name_val]
                has_tasks = any(ms_render_data.get('issue_types', {}).values()) or ms_render_data.get(
                    'tasks_without_type_grouping')
                if not has_tasks:
                    logger.debug(f"МС '{ms_name_val}' в '{section_id}' не содержит задач для Word.")
                    continue
                _add_heading_styled(document_obj, ms_name_val, level=3, style_to_apply=style_h3_ms)

                group_by_type_flag = current_section_data.get('group_by_issue_type', False)
                render_queue: List[Dict] = []
                if group_by_type_flag:
                    issue_types_data = ms_render_data.get('issue_types', {})
                    for type_name_str in sorted(issue_types_data.keys()):
                        tasks = issue_types_data[type_name_str]
                        if tasks:
                            render_queue.append(
                                {"is_header": True, "text": type_name_str, "level_style_name": style_h4_type})
                            render_queue.extend([{"is_header": False, "data": t_dict} for t_dict in
                                                 sorted(tasks, key=lambda t: str(t.get("key", "")))])
                else:
                    tasks_no_group = ms_render_data.get('tasks_without_type_grouping', [])
                    if tasks_no_group:
                        render_queue.extend([{"is_header": False, "data": t_dict} for t_dict in
                                             sorted(tasks_no_group, key=lambda t: str(t.get("key", "")))])

                if not render_queue: continue
                for item in render_queue:
                    if item.get("is_header"):
                        _add_heading_styled(document_obj, item["text"], level=4,
                                            style_to_apply=item["level_style_name"])
                    else:
                        _add_task_entry_to_document(document_obj, issue_template, item["data"],
                                                    style_task_first, style_task_multiline, multiline_indent_val)
                document_obj.add_paragraph()

    logger.info("Генерация Word документа успешно завершена.")
    return document_obj